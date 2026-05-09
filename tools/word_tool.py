"""
Appends formatted content to a Word (.docx) document.
Parses ChatGPT's markdown output and applies the styles defined in config.json.
"""

import json
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .parser import parse_markdown, Block, Run

PENDING_SUFFIX = ".pending.json"

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
CONFIG_PATH = os.path.join(ROOT, "config.json")

ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

STYLE_KEY_MAP = {
    "chapter":            "chapter",
    "heading2":           "heading2",
    "heading3":           "heading3",
    "body":               "body",
    "list_item":          "body",
    "artifact":           "body",
    "table":              "body",
    "table_caption":      "body",
    "equation":           "body",
    "figure":             "body",
    "figure_caption":     "body",
    "figure_placeholder": "body",
}

WORD_STYLE_MAP = {
    "chapter":            "Heading 1",
    "heading2":           "Heading 2",
    "heading3":           "Heading 3",
    "body":               "Normal",
    "list_item":          "Normal",
    "artifact":           "Normal",
    "table":              "Normal",
    "table_caption":      "Normal",
    "equation":           "Normal",
    "figure":             "Normal",
    "figure_caption":     "Normal",
    "figure_placeholder": "Normal",
}


def _load_styles() -> dict:
    with open(CONFIG_PATH, encoding="utf-8") as f:
        return json.load(f).get("styles", {})


def _apply_paragraph_format(para, cfg: dict):
    pf = para.paragraph_format
    if "alignment" in cfg:
        pf.alignment = ALIGN_MAP.get(cfg["alignment"], WD_ALIGN_PARAGRAPH.LEFT)
    if "space_before" in cfg:
        pf.space_before = Pt(cfg["space_before"])
    if "space_after" in cfg:
        pf.space_after = Pt(cfg["space_after"])
    if "first_line_cm" in cfg:
        pf.first_line_indent = Cm(cfg["first_line_cm"])
    if "hanging_cm" in cfg:
        pf.left_indent = Cm(cfg["hanging_cm"])
        pf.first_line_indent = Cm(-cfg["hanging_cm"])


def _apply_run_format(run, cfg: dict, bold: bool = False, italic: bool = False):
    run.font.name = cfg.get("font", "Garamond")
    run.font.size = Pt(cfg.get("size_pt", 11))
    run.bold = cfg.get("bold", False) or bold
    run.italic = italic
    if cfg.get("all_caps", False):
        run.font.all_caps = True


def _configure_page(doc: Document, page_cfg: dict):
    section = doc.sections[0]
    if "width_cm" in page_cfg:
        section.page_width = Cm(page_cfg["width_cm"])
    if "height_cm" in page_cfg:
        section.page_height = Cm(page_cfg["height_cm"])
    if "margin_cm" in page_cfg:
        m = Cm(page_cfg["margin_cm"])
        section.left_margin = section.right_margin = m
        section.top_margin = section.bottom_margin = m
    if "header_dist_cm" in page_cfg:
        section.header_distance = Cm(page_cfg["header_dist_cm"])
    if "footer_dist_cm" in page_cfg:
        section.footer_distance = Cm(page_cfg["footer_dist_cm"])


def _open_document(word_file_path: str, styles: dict) -> Document:
    if os.path.exists(word_file_path):
        return Document(word_file_path)

    doc = Document()
    page_cfg = styles.get("page", {})
    if page_cfg:
        _configure_page(doc, page_cfg)
    return doc


_PENDING_EQ_PREFIX = "PENDING_EQ:"

# ---------------------------------------------------------------------------
# OpenAI equation cleaning
# ---------------------------------------------------------------------------

_OPENAI_EQ_SYSTEM_PROMPT = (
    "You are a Microsoft Word equation formatter. "
    "Convert the input equation — which may be garbled, in LaTeX, or rendered "
    "Unicode — into clean Microsoft Word UnicodeMath (Linear Format) as used in "
    "the Word equation editor (Alt+=).\n\n"
    "Rules:\n"
    "- Subscripts with _, superscripts with ^.  Group with () when needed.\n"
    "- Greek: use Unicode directly (α β γ δ ε ζ η θ λ μ ν π σ τ φ ψ ω etc.).\n"
    "- Sums/products: ∑_(i=1)^n x_i  ∏_(k=1)^K p_k\n"
    "- Integrals: ∫_a^b f(x)ⅆx\n"
    "- Fractions: a/b or (a+b)/(c+d); complex: \\frac(num)(den)\n"
    "- Indicator 𝟙[cond] or mathbb1[cond]: write 𝟙[condition]\n"
    "- argmax/argmin: keep as text with subscript, e.g. argmax_(θ)\n"
    "- land/lor → ∧/∨;  le/ge → ≤/≥;  in → ∈\n"
    "- mathbb{R}/ℝ→ℝ  mathbb{N}→ℕ  mathbb{E}/E→𝔼  mathbb{1}/1→𝟙\n"
    "- tau→τ  lambda→λ  epsilon→ε  theta→θ  alpha→α etc.\n"
    "- Remove \\tag{n}, tag{n}, tagN, \\label{...}, \\nonumber, $ signs.\n"
    "- Remove \\mathrm, \\text, \\mathbf etc. wrappers; keep their content.\n"
    "Return ONLY the clean UnicodeMath equation. No markdown. No explanations."
)


def _openai_clean_equation(text: str) -> "str | None":
    """Call OpenAI to convert a garbled/Unicode equation to Word UnicodeMath."""
    api_key = os.environ.get("OPENAI_API_KEY", "")
    if not api_key:
        return None
    try:
        import openai  # noqa: PLC0415
        client = openai.OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model=os.environ.get("OPENAI_EQ_MODEL", "gpt-4o-mini"),
            temperature=0,
            max_tokens=512,
            messages=[
                {"role": "system", "content": _OPENAI_EQ_SYSTEM_PROMPT},
                {"role": "user", "content": text},
            ],
        )
        result = resp.choices[0].message.content
        if result:
            return result.strip()
    except Exception as exc:
        print(f"  [OpenAI] equation conversion failed: {exc}")
    return None


# ---------------------------------------------------------------------------

_OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_NARY_OPS = {"∑", "∏", "∫", "∬", "∭", "∮", "⋃", "⋂"}


def _mml_to_omml_node(mml, parent):
    """Recursively convert one MathML element into OMML children of *parent*."""
    from lxml import etree

    ns = _OMML_NS
    tag = mml.tag.split("}")[-1] if "}" in mml.tag else mml.tag
    ch = list(mml)

    def sub(tag_name, elem=parent):
        return etree.SubElement(elem, f"{{{ns}}}{tag_name}")

    def run(text, italic=False, p=parent):
        r = etree.SubElement(p, f"{{{ns}}}r")
        if italic:
            rpr = etree.SubElement(r, f"{{{ns}}}rPr")
            sty = etree.SubElement(rpr, f"{{{ns}}}sty")
            sty.set(f"{{{ns}}}val", "i")
        t = etree.SubElement(r, f"{{{ns}}}t")
        t.text = text
        return r

    def recurse(mml_node, omml_parent):
        _mml_to_omml_node(mml_node, omml_parent)

    # Transparent containers
    if tag in ("math", "mrow", "mstyle", "merror", "mpadded",
               "mphantom", "semantics", "mfenced"):
        for child in ch:
            recurse(child, parent)

    elif tag == "annotation":
        pass  # skip LaTeX source annotation

    elif tag in ("mi", "mn", "mo", "mtext", "ms"):
        text = "".join(mml.itertext())
        if text.startswith("\\"):
            text = text[1:]  # strip backslash — latex2mathml leaves \arg, \max, etc.
        if text:
            run(text, italic=(tag == "mi" and len(text) == 1))

    elif tag == "msub" and len(ch) >= 2:
        ssub = sub("sSub")
        recurse(ch[0], sub("e", ssub))
        recurse(ch[1], sub("sub", ssub))

    elif tag == "msup" and len(ch) >= 2:
        ssup = sub("sSup")
        recurse(ch[0], sub("e", ssup))
        recurse(ch[1], sub("sup", ssup))

    elif tag == "msubsup" and len(ch) >= 3:
        sss = sub("sSubSup")
        recurse(ch[0], sub("e", sss))
        recurse(ch[1], sub("sub", sss))
        recurse(ch[2], sub("sup", sss))

    elif tag == "mfrac" and len(ch) >= 2:
        f = sub("f")
        recurse(ch[0], sub("num", f))
        recurse(ch[1], sub("den", f))

    elif tag == "msqrt":
        rad = sub("rad")
        radpr = sub("radPr", rad)
        dh = sub("degHide", radpr)
        dh.set(f"{{{ns}}}val", "1")
        sub("deg", rad)  # empty = square root
        e = sub("e", rad)
        for child in ch:
            recurse(child, e)

    elif tag == "mroot" and len(ch) >= 2:
        rad = sub("rad")
        recurse(ch[1], sub("deg", rad))
        recurse(ch[0], sub("e", rad))

    elif tag in ("munder", "mover", "munderover"):
        op_text = "".join(ch[0].itertext()).strip() if ch else ""
        if op_text in _NARY_OPS:
            nary = sub("nary")
            narypr = sub("naryPr", nary)
            chr_e = sub("chr", narypr)
            chr_e.set(f"{{{ns}}}val", op_text)
            if tag == "munder":
                h = sub("supHide", narypr); h.set(f"{{{ns}}}val", "1")
                recurse(ch[1], sub("sub", nary))
                sub("sup", nary)
            elif tag == "mover":
                h = sub("subHide", narypr); h.set(f"{{{ns}}}val", "1")
                sub("sub", nary)
                recurse(ch[1], sub("sup", nary))
            else:
                recurse(ch[1], sub("sub", nary))
                recurse(ch[2], sub("sup", nary))
            sub("e", nary)
        elif tag == "munder":
            ll = sub("limLow")
            recurse(ch[0], sub("e", ll))
            if len(ch) >= 2: recurse(ch[1], sub("lim", ll))
        elif tag == "mover":
            lu = sub("limUpp")
            recurse(ch[0], sub("e", lu))
            if len(ch) >= 2: recurse(ch[1], sub("lim", lu))
        else:  # munderover non-nary
            sss = sub("sSubSup")
            recurse(ch[0], sub("e", sss))
            if len(ch) >= 2: recurse(ch[1], sub("sub", sss))
            if len(ch) >= 3: recurse(ch[2], sub("sup", sss))

    elif tag == "mtable":
        m = sub("m")
        for mtr in ch:
            mr = sub("mr", m)
            for mtd in mtr:
                e = sub("e", mr)
                for child in mtd:
                    recurse(child, e)

    else:
        text = "".join(mml.itertext())
        if text.strip():
            run(text)


_LATEX_STRIP_RE = re.compile(
    r"\\tag\*?\{[^}]*\}"     # \tag{11} / \tag*{11}
    r"|\\label\{[^}]*\}"     # \label{eq:foo}
    r"|\\notag\b"             # \notag
    r"|\\nonumber\b",         # \nonumber
    re.IGNORECASE,
)


def _latex_to_omml(latex: str):
    """
    Convert a LaTeX string to an OMML lxml element (no XSLT file needed).
    Uses latex2mathml for LaTeX→MathML, then a Python MathML→OMML converter.
    Returns None on any failure.
    """
    try:
        import latex2mathml.converter
        from lxml import etree
    except ImportError:
        return None

    # Strip equation-numbering commands — Word doesn't use LaTeX numbering
    latex = _LATEX_STRIP_RE.sub("", latex).strip()

    try:
        mathml_str = latex2mathml.converter.convert(latex)
        mathml_tree = etree.fromstring(mathml_str.encode("utf-8"))
    except Exception:
        return None

    try:
        omath = etree.Element(f"{{{_OMML_NS}}}oMath")
        _mml_to_omml_node(mathml_tree, omath)
        return omath
    except Exception:
        return None


def _write_equation_block(doc: Document, block, styles: dict):
    """
    Write a display equation into the Word document.

    Priority:
    1. OMML via latex2mathml (best — proper Word equation object, no COM needed)
    2. OpenAI → UnicodeMath PENDING_EQ: marker  (converted by _update_word_fields via COM)
    3. Centered italic raw text fallback (readable but not an equation box)
    """
    latex = block.latex
    if not latex:
        return

    cfg = styles.get("body", {})

    omml = _latex_to_omml(latex)
    if omml is not None:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para._p.append(omml)
        return

    # Try OpenAI → UnicodeMath; write PENDING_EQ: marker for COM to finish
    linear_math = _openai_clean_equation(latex)
    if linear_math:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"{_PENDING_EQ_PREFIX}{linear_math}")
        run.font.name = cfg.get("font", "Garamond")
        run.font.size = Pt(cfg.get("size_pt", 11))
        return

    # Final fallback: centered italic text so content is not silently dropped
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(latex)
    run.font.name = cfg.get("font", "Garamond")
    run.font.size = Pt(cfg.get("size_pt", 11))
    run.italic = True


def _write_table_block(doc: Document, block, styles: dict):
    """Add a Word table for a 'table' block. First row is treated as header."""
    rows = block.table_rows
    if not rows:
        return
    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    cfg = styles.get("body", {})
    font_name = cfg.get("font", "Garamond")
    font_size = Pt(cfg.get("size_pt", 11))

    for r_idx, row_data in enumerate(rows):
        is_header = r_idx == 0
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(cell_text)
            run.font.name = font_name
            run.font.size = font_size
            run.bold = is_header

        # Fill any missing cells in this row
        for c_idx in range(len(row_data), num_cols):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run("")
            run.font.name = font_name
            run.font.size = font_size

    # Blank paragraph after table — caption (if any) will follow this
    doc.add_paragraph()


def _write_figure_block(doc: Document, block, styles: dict):
    """Insert a downloaded figure image centered with a configurable max width."""
    cfg = styles.get("body", {})
    img_path = block.figure_image_path

    if img_path and os.path.exists(img_path):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run()
        try:
            run.add_picture(img_path, width=Cm(14))
        except Exception:
            run.text = f"[Figure {block.figure_number} — image load error]"
            run.font.italic = True
    else:
        # No image yet — write a visible placeholder so the spot is reserved
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[Figure {block.figure_number}]")
        run.font.name = cfg.get("font", "Garamond")
        run.font.size = Pt(cfg.get("size_pt", 11))
        run.italic = True


_FIGURE_LABEL_PREFIX_RE = re.compile(
    r"^\s*(?:fig\.?|figure)\s+\d+\s*[.:\-–—]?\s*", re.I
)
_TABLE_LABEL_PREFIX_RE = re.compile(
    r"^\s*table\s+\d+\s*[.:\-–—]?\s*", re.I
)


def _write_word_caption(doc: Document, label: str, caption_text: str = ""):
    """
    Insert a Word auto-numbered caption using the built-in Caption style.
    Produces the same output as Insert → Caption in Word.
    label: "Figure" or "Table"
    caption_text: descriptive text that follows the auto-number
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        para = doc.add_paragraph(style="Caption")
    except Exception:
        para = doc.add_paragraph()

    p = para._p

    def _run(text):
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        return r

    def _fldchar(fld_type):
        r = OxmlElement("w:r")
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), fld_type)
        r.append(fc)
        return r

    # "Figure " / "Table " label
    p.append(_run(f"{label} "))

    # SEQ field  →  SEQ Figure \* ARABIC  or  SEQ Table \* ARABIC
    p.append(_fldchar("begin"))
    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {label} \\* ARABIC "
    r_instr.append(instr)
    p.append(r_instr)
    p.append(_fldchar("separate"))
    p.append(_run("1"))          # cached placeholder — Word updates on open
    p.append(_fldchar("end"))

    # Descriptive caption text
    if caption_text.strip():
        p.append(_run(f" {caption_text.strip()}"))


def _strip_caption_prefix(text: str, label: str) -> str:
    """Remove leading label+number from caption text (Word adds the number itself)."""
    if label == "Figure":
        return _FIGURE_LABEL_PREFIX_RE.sub("", text).strip()
    if label == "Table":
        return _TABLE_LABEL_PREFIX_RE.sub("", text).strip()
    return text.strip()


def _write_figure_caption_block(doc: Document, block, styles: dict):
    """Write a figure caption using Word's built-in Caption style + SEQ field."""
    caption_text = _strip_caption_prefix(block.plain_text, "Figure")
    _write_word_caption(doc, "Figure", caption_text)


def _convert_pending_equations(word_doc) -> int:
    """
    Find PENDING_EQ: paragraphs in a live Word COM Document and convert each
    to a proper Word OMath equation object via OMaths.Add().
    Returns the count of equations converted.
    """
    prefix = _PENDING_EQ_PREFIX
    prefix_len = len(prefix)
    wdCharacter = 1  # Word VBA constant for character movement unit

    # Collect indices first (reverse iteration avoids index shifts)
    pending_indices = []
    total = word_doc.Paragraphs.Count
    for i in range(1, total + 1):
        try:
            if word_doc.Paragraphs(i).Range.Text.startswith(prefix):
                pending_indices.append(i)
        except Exception:
            pass

    converted = 0
    for i in reversed(pending_indices):
        try:
            para = word_doc.Paragraphs(i)
            full_text = para.Range.Text
            linear_math = full_text[prefix_len:].rstrip("\r\n\x07")
            if not linear_math.strip():
                continue
            rng = para.Range
            # Exclude the paragraph mark from the range before setting text
            rng.MoveEnd(wdCharacter, -1)
            rng.Text = linear_math
            rng.ParagraphFormat.Alignment = 1  # wdAlignParagraphCenter
            word_doc.OMaths.Add(rng)
            converted += 1
        except Exception as exc:
            print(f"  [Word] PENDING_EQ conversion failed at para {i}: {exc}")

    return converted


def _update_word_fields(word_file_path: str) -> None:
    """Recalculate all SEQ / caption fields and convert PENDING_EQ equation
    markers to real Word equation objects via COM automation.

    python-docx writes SEQ fields with a cached placeholder of "1" and writes
    PENDING_EQ: markers for equations that need Word's OMaths.Add().
    This function opens (or reuses) the live Word instance, fixes both, and
    saves.  Silently no-ops when pywin32 or Word is unavailable.
    """
    try:
        import win32com.client  # noqa: PLC0415
    except ImportError:
        return  # pywin32 not installed

    abs_path = os.path.abspath(word_file_path)
    word = None
    word_was_running = True
    opened_here = False

    try:
        # Prefer an already-running Word instance to avoid spawning extras
        try:
            word = win32com.client.GetActiveObject("Word.Application")
        except Exception:
            word = win32com.client.Dispatch("Word.Application")
            word_was_running = False
            word.Visible = False

        # Check whether this document is already open in Word
        doc = None
        try:
            for d in word.Documents:
                if os.path.normcase(d.FullName) == os.path.normcase(abs_path):
                    doc = d
                    break
        except Exception:
            pass

        if doc is None:
            doc = word.Documents.Open(abs_path)
            opened_here = True

        doc.Fields.Update()

        converted = _convert_pending_equations(doc)
        if converted:
            print(f"  [Word] converted {converted} pending equation(s) to OMath objects")

        doc.Save()

        if opened_here:
            doc.Close(SaveChanges=False)

    except Exception as exc:
        print(f"  [Word] COM update skipped: {exc}")
    finally:
        if word is not None and not word_was_running:
            try:
                word.Quit()
            except Exception:
                pass


def append_blocks_to_word(blocks: list[Block], word_file_path: str) -> str:
    """Append pre-parsed blocks to the Word document using configured styles."""
    styles = _load_styles()
    doc = _open_document(word_file_path, styles)

    if not blocks:
        return "Warning: no content blocks parsed from response."

    # table_caption blocks appear before their table in the parsed block list
    # but must be emitted AFTER the table in Word (below it).
    pending_table_caption: Block | None = None

    for block in blocks:
        if block.block_type == "table_caption":
            # Hold it — write after the next table
            pending_table_caption = block
            continue

        if block.block_type == "table":
            _write_table_block(doc, block, styles)
            cap_text = _strip_caption_prefix(
                pending_table_caption.plain_text if pending_table_caption else "", "Table"
            )
            _write_word_caption(doc, "Table", cap_text)
            pending_table_caption = None
            continue

        if pending_table_caption:
            # A table_caption appeared but was not immediately followed by a table;
            # flush it as a plain caption before continuing.
            _write_word_caption(doc, "Table", pending_table_caption.plain_text)
            pending_table_caption = None

        # Artifacts are placement/description notes — skip in final Word output
        if block.block_type == "artifact":
            continue

        if block.block_type == "equation":
            _write_equation_block(doc, block, styles)
            continue

        if block.block_type in ("figure", "figure_placeholder"):
            _write_figure_block(doc, block, styles)
            continue

        if block.block_type == "figure_caption":
            _write_figure_caption_block(doc, block, styles)
            continue

        style_key = STYLE_KEY_MAP.get(block.block_type, "body")
        word_style = WORD_STYLE_MAP.get(block.block_type, "Normal")
        cfg = styles.get(style_key, {})

        # Add the paragraph using the built-in Word style
        try:
            para = doc.add_paragraph(style=word_style)
        except KeyError:
            para = doc.add_paragraph()

        _apply_paragraph_format(para, cfg)

        # Bullet prefix for list items
        prefix = "• " if block.block_type == "list_item" else ""
        first_run = True
        for run_data in block.runs:
            text = (prefix + run_data.text) if first_run and prefix else run_data.text
            first_run = False
            if not text:
                continue
            run = para.add_run(text)
            _apply_run_format(run, cfg, bold=run_data.bold, italic=run_data.italic)

    # Flush any trailing table caption that had no table after it
    if pending_table_caption:
        _write_word_caption(doc, "Table", pending_table_caption.plain_text)

    doc.save(word_file_path)
    _update_word_fields(word_file_path)
    return f"OK: wrote {len(blocks)} block(s) to '{word_file_path}'"


def append_to_word(content: str, word_file_path: str) -> str:
    """
    Parse *content* and append it to the Word document at *word_file_path*,
    applying the configured styles. Creates the file if it does not exist.
    """
    return append_blocks_to_word(parse_markdown(content), word_file_path)


# ── Pending-file helpers ──────────────────────────────────────────────────────

def _pending_path(word_file_path: str) -> str:
    return word_file_path + PENDING_SUFFIX


def _blocks_to_json(blocks: list[Block]) -> list[dict]:
    return [
        {
            "block_type": b.block_type,
            "runs": [{"text": r.text, "bold": r.bold, "italic": r.italic} for r in b.runs],
            "table_rows": b.table_rows,
            "latex": b.latex,
            "figure_number": b.figure_number,
            "figure_image_path": b.figure_image_path,
        }
        for b in blocks
    ]


def _blocks_from_json(data: list[dict]) -> list[Block]:
    result = []
    for entry in data:
        runs = [Run(text=r["text"], bold=r["bold"], italic=r["italic"]) for r in entry["runs"]]
        result.append(Block(
            block_type=entry["block_type"],
            runs=runs,
            table_rows=entry.get("table_rows", []),
            latex=entry.get("latex", ""),
            figure_number=entry.get("figure_number", 0),
            figure_image_path=entry.get("figure_image_path", ""),
        ))
    return result


def save_pending_blocks(blocks: list[Block], word_file_path: str):
    """Append *blocks* to the pending file for *word_file_path*."""
    path = _pending_path(word_file_path)
    existing: list[dict] = []
    if os.path.exists(path):
        try:
            with open(path, encoding="utf-8") as f:
                existing = json.load(f)
        except Exception:
            existing = []
    existing.extend(_blocks_to_json(blocks))
    with open(path, "w", encoding="utf-8") as f:
        json.dump(existing, f, ensure_ascii=False, indent=2)


def has_pending_blocks(word_file_path: str) -> bool:
    return os.path.exists(_pending_path(word_file_path))


def recover_pending_blocks(word_file_path: str) -> str:
    """
    Read the pending file for *word_file_path*, append its blocks to the
    Word document, and delete the pending file on success.
    """
    path = _pending_path(word_file_path)
    if not os.path.exists(path):
        return "No pending content found for this document."

    with open(path, encoding="utf-8") as f:
        raw = json.load(f)

    blocks = _blocks_from_json(raw)
    if not blocks:
        os.remove(path)
        return "Pending file was empty — nothing to recover."

    result = append_blocks_to_word(blocks, word_file_path)
    os.remove(path)
    return f"Recovered {len(blocks)} block(s). {result}"
